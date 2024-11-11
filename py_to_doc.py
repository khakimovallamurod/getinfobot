from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def creat_docfile():
    # Hujjat yaratish
    doc = Document()

    # Sahifa cheklovlarini o'rnatish
    section = doc.sections[0]
    section.top_margin = Inches(0.59)   # Yuqoridan 1.5 sm (0.59 dyuym)
    section.bottom_margin = Inches(0.39)  # Pastdan 1 sm (0.39 dyuym)
    section.left_margin = Inches(0.79)    # Chapdan 2 sm (0.79 dyuym)
    section.right_margin = Inches(0.39)   # O'ngdan 1 sm (0.39 dyuym)

    # Matnni markazlashtirib kiritish
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("MA'LUMOTNOMA")
    run.bold = True
    run.font.size = Pt(14)
    return doc


def add_fio(doc, fullname: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(fullname)
    run.bold = True
    run.font.size = Pt(14)

def user_data(doc, work_name: str, left_data: dict, right_data: dict, button_data: dict, image_path: str):
    user_table = doc.add_table(rows=1, cols=2)
    user_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Avtomatik kenglikni moslashtirish
    user_table.autofit = True

    # Set custom column widths
    col1, col2 = user_table.rows[0].cells
    col1.width = Inches(12)  # Birinchi ustun ko'proq joy oladi
    col2.width = Inches(1.18)  # Ikkinchi ustun faqat 3x4 sm sig'adigan o'lchamda bo'ladi

    paragraph = col1.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(work_name)

    table_data = col1.add_table(rows=1, cols=2)
    table_data.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell1, cell2 = table_data.rows[0].cells

    # left write
    paragraph = cell1.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx, (key, val) in enumerate(left_data.items()):
        if idx==len(left_data)-1:
            run = paragraph.add_run(f"{key}: \n")
            run.bold = True
            paragraph.add_run(val)
        else:
            run = paragraph.add_run(f"{key}: \n")
            run.bold = True
            paragraph.add_run(val+'\n')

    # right write
    paragraph = cell2.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    #1
    for idx, (key, val) in enumerate(right_data.items()):
        if idx==len(right_data)-1:
            run = paragraph.add_run(f"{key}: \n")
            run.bold = True
            paragraph.add_run(val)
        else:
            run = paragraph.add_run(f"{key}: \n")
            run.bold = True
            paragraph.add_run(val+'\n')

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for idx, (key, val) in enumerate(button_data.items()):
        if idx==len(button_data)-1:
            run = paragraph.add_run(f"{key}: \n")
            run.bold = True
            paragraph.add_run(val)
        else:
            run = paragraph.add_run(f"{key}: \n")
            run.bold = True
            paragraph.add_run(val+'\n')
    # Rasmni 3x4 sm qilib joylashtirish
    paragraph = col2.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(1.18), height=Inches(1.575))  # Rasmni 3x4 sm o'lchamiga moslashtirish


def mehnat_faoliyati(doc, text: str):
    # Sarlavhani qo'shish
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("MEHNAT FAOLIYATI")
    run.bold = True
    run.font.size = Pt(14)

    # Matnni qatorma-qator qo'shish
    for line in text.splitlines():
        paragraph = doc.add_paragraph(line)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(6)  # Qatorlar orasiga bo'sh joy qo'shish


def yaqin_qarindosh_malumoti(doc, fullname: str, qarindosh_data: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"{fullname.title()}ning yaqin qarindoshlari xaqida\nMA'LUMOTNOMA")
    run.bold = True
    run.font.size = Pt(14)

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # Sarlavhalarni qo'shish
    headers = ["Qarindoshligi", "Familyasi, ismi va otasining ismi", "Tug'ilgan yili va joyi", "Ish joyi va lavozimi", "Turar joyi"]
    hdr_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    # Ma'lumotlarni qo'shish
    for line in qarindosh_data.splitlines():
        data = line.split(':')
        row_cells = table.add_row().cells
        row_cells[0].text = data[0]
        row_cells[1].text = data[1]
        row_cells[2].text = data[2]
        row_cells[3].text = data[3]
        row_cells[4].text = data[4]

        # Matnni chiroyli ko'rinishda formatlash
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)


def save_obektvka(doc, file_name):
    doc.save(file_name)

def control(data_all: str):
    doc = creat_docfile()
    name = data_all['name']
    #1
    add_fio(doc, name)
    #2
    work_name = data_all['work']
    left_data = {
        "Tu'gilgan yili": data_all['year'],
        "Millati": data_all['millati'],
        "Malumoti": data_all['malumoti'],
        "Ma'lumoti bo'yicha mutaxassisligi": data_all['mutaxasis'],
        "Ilmiy darajasi": data_all['rank'],
        "Qaysi chet tillarni biladi": data_all['chettili'],
    }

    right_data = {
        "Tug'ilgan joyi": data_all['loaction'],
        "Partiyaviyligi": data_all['partiyaviy'],
        "Tamomlagan": data_all['tamomlagan'],
        "Ilmiy unvoni": data_all['unvon'],
        "Xarbiy (maxsus) unvoni": data_all['xunvon']
    }
    button_data = {
        "Davlat mukofatlari bilan taqdirlanganmi (qanaqa)": data_all['medal'],
        "Xalq deputatlari, respublika, viloyat, shahar va tuman Kengashi deputatimi yoki boshqa saylangan organlarning azosimi (to'liq ko'rsatilishi lozim)": data_all['yuqorimansab']
    }
    image_path = data_all['image']

    user_data(doc, work_name, left_data=left_data, right_data=right_data, button_data=button_data, image_path=image_path)
    #3
    mehnatfaoliyat = data_all['mehnatfaoliyat']
    mehnat_faoliyati(doc, text=mehnatfaoliyat)
    #4 
    yaqinqarindosh = data_all['yaqinqarindosh']
    yaqin_qarindosh_malumoti(doc, fullname=name, qarindosh_data = yaqinqarindosh)
    #5
    save_obektvka(doc, f'{name}.docx')

